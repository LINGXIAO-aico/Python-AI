"""文档加载与抽取测试。"""

from __future__ import annotations

import json

import pytest
from docx import Document

from campus_rag.loaders import clean_text, export_extracted_jsonl, extract_text


def test_clean_text_removes_extra_spaces_and_page_marks() -> None:
    text = clean_text("  第一行   内容\n\n\n第 12 页\n第二行\t内容  ")

    assert text == "第一行 内容\n\n\n第二行 内容"


def test_extract_text_from_txt_and_md(tmp_path) -> None:
    txt = tmp_path / "notice.txt"
    md = tmp_path / "guide.md"
    txt.write_text("校园卡   补办\n第 1 页", encoding="utf-8")
    md.write_text("# 图书馆\n\n开放时间", encoding="utf-8")

    assert extract_text(txt) == "校园卡 补办"
    assert "图书馆" in extract_text(md)


def test_extract_text_from_docx(tmp_path) -> None:
    path = tmp_path / "notice.docx"
    doc = Document()
    doc.add_paragraph("学生证补办流程")
    doc.add_paragraph("")
    doc.add_paragraph("携带身份证到学生事务中心。")
    doc.save(path)

    text = extract_text(path)

    assert "学生证补办流程" in text
    assert "学生事务中心" in text


def test_extract_text_rejects_unsupported_suffix(tmp_path) -> None:
    path = tmp_path / "data.csv"
    path.write_text("x", encoding="utf-8")

    with pytest.raises(ValueError, match="暂不支持"):
        extract_text(path)


def test_export_extracted_jsonl_skips_empty_files(tmp_path) -> None:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    (source_dir / "a.txt").write_text("选课通知", encoding="utf-8")
    (source_dir / "b.md").write_text("   ", encoding="utf-8")
    output_path = tmp_path / "out" / "extracted.jsonl"

    count = export_extracted_jsonl(source_dir, output_path)
    rows = [json.loads(line) for line in output_path.read_text(encoding="utf-8").splitlines()]

    assert count == 1
    assert rows[0]["doc_id"] == "SRC001"
    assert rows[0]["title"] == "a"
