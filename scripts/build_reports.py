from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
FIGURE_DIR = REPORT_DIR / "figures"
LOG_DIR = ROOT / "logs"
REPORT_DATE = "2026年5月10日"
TITLE = "B07 RAG 校园智能问答助手"
TEAM_MEMBERS = "归梦依、凌霄、周子涵、杨歆苒"
TEAM_ASSIGNMENT = "归梦依：数据处理与知识库；凌霄：检索策略设计与优化；周子涵：大模型调用与系统开发；杨歆苒：评测分析与报告撰写"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


DATA_PROFILE = read_json(LOG_DIR / "data_profile.json")
TRAINING_LOG = read_json(LOG_DIR / "training_log.json")
EVAL_SUMMARY = read_json(LOG_DIR / "evaluation_summary.json")
CATEGORY_NAMES = "、".join(DATA_PROFILE["categories"].keys())
CATEGORY_SUMMARY = f"覆盖{len(DATA_PROFILE['categories'])}类校园场景：{CATEGORY_NAMES}。"
QUESTION_COUNT = int(EVAL_SUMMARY["question_count"])


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", size=10.5, bold=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=13 if level == 1 else 11.5, bold=True, color=(31, 78, 121))
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table, header=True):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, bold=(header and row_idx == 0))
            if header and row_idx == 0:
                shade_cell(cell, "D9EAF7")


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    return doc


def build_selection_docx(path: Path) -> None:
    doc = setup_doc()
    add_paragraph(doc, "《Python 人工智能程序设计实践》选题报告", 16, True, (31, 78, 121), WD_ALIGN_PARAGRAPH.CENTER, 2)
    add_paragraph(doc, TITLE, 14, True, None, WD_ALIGN_PARAGRAPH.CENTER, 8)

    meta = doc.add_table(rows=4, cols=4)
    values = [
        ("项目编号", "B07", "提交日期", REPORT_DATE),
        ("项目名称", "RAG 校园智能问答助手", "题目来源", "课程项目库 B 类"),
        ("小组成员", TEAM_MEMBERS, "组队人数", "4人"),
        ("团队分工", TEAM_ASSIGNMENT, "代码仓库", "本地已建，可上传 GitHub/Gitee"),
    ]
    for row, row_values in zip(meta.rows, values):
        for idx, text in enumerate(row_values):
            row.cells[idx].text = text
    style_table(meta, header=False)

    add_heading(doc, "一、项目简介")
    add_paragraph(
        doc,
        "本项目拟构建面向校园办事、学习生活与公共服务场景的智能问答助手。系统不训练大模型，"
        "而是采用检索增强生成（RAG）范式：先从校园知识库中召回相关文本块，再基于证据生成带来源引用的回答，"
        "解决通用模型容易回答不准、无法说明依据的问题，并形成 RAG 系统与无检索基线、不同检索策略之间的对比实验。",
        9.9,
    )
    add_heading(doc, "二、数据来源与合规性")
    add_paragraph(
        doc,
        f"现阶段已整理模拟校园 FAQ 与办事指南 {DATA_PROFILE['clean_rows']} 条，覆盖 {len(DATA_PROFILE['categories'])} 类场景；"
        "数据字段包括编号、类别、标题、正文、来源、链接和更新时间。后续由归梦依补充学校官网学生手册、教务规定、图书馆与宿舍办事指南等 PDF/Word 原始资料；"
        "当前数据均为课程演示用公开/自建文本，不含个人隐私信息。",
        9.9,
    )
    add_heading(doc, "三、技术路线")
    route = doc.add_table(rows=1, cols=4)
    route.rows[0].cells[0].text = "文档抽取与清洗"
    route.rows[0].cells[1].text = "切分、向量化与建库"
    route.rows[0].cells[2].text = "BM25/向量/混合检索"
    route.rows[0].cells[3].text = "RAG回答、评测与Demo"
    style_table(route, header=False)
    add_paragraph(
        doc,
        "基线实现采用 pdfplumber/python-docx 抽取文档、Pandas 清洗、Scikit-learn TF-IDF 建立向量索引，并实现自定义 BM25 与 0.5/0.5 混合检索。"
        "回答模块预留通义千问/OpenAI 兼容 API；无 Key 时使用离线抽取式回答保证可复现。Demo 提供 CLI、Streamlit 和 Gradio 三种入口。",
        9.9,
    )
    add_heading(doc, "四、参考代码与资料")
    add_paragraph(
        doc,
        "参考 scikit-learn TfidfVectorizer、Streamlit、Gradio、OpenAI Python SDK、DashScope 兼容接口文档，以及 RAG 相关论文和开源实现；"
        "实际代码已在本项目目录中独立实现，后续上传仓库时将在 README 中保留引用说明。",
        9.9,
    )
    add_heading(doc, "五、进度计划")
    plan = doc.add_table(rows=4, cols=3)
    plan_rows = [
        ("第7周", "完成选题、数据源确定、项目结构搭建", "选题报告"),
        ("第8-10周", "完成数据清洗、基线检索、BM25/混合检索、评测问题集", "中期进展报告"),
        ("第11-14周", "扩充真实知识库到100条以上，引入 Embedding/FAISS 或大模型 API，完善可视化", "代码仓库更新"),
        ("第15-16周", "整理技术报告、PPT 与演示 Demo，完成答辩", "最终材料"),
    ]
    for row, values in zip(plan.rows, plan_rows):
        for idx, text in enumerate(values):
            row.cells[idx].text = text
    style_table(plan, header=False)
    doc.save(path)


def build_midterm_docx(path: Path) -> None:
    doc = setup_doc()
    add_paragraph(doc, "《Python 人工智能程序设计实践》中期进展报告", 16, True, (31, 78, 121), WD_ALIGN_PARAGRAPH.CENTER, 2)
    add_paragraph(doc, TITLE, 14, True, None, WD_ALIGN_PARAGRAPH.CENTER, 6)

    meta = doc.add_table(rows=3, cols=4)
    values = [
        ("项目编号", "B07", "提交日期", REPORT_DATE),
        ("小组成员", TEAM_MEMBERS, "当前阶段", "已跑通基线模型"),
        ("代码目录", "当前项目目录", "Demo", "Streamlit + CLI"),
    ]
    for row, row_values in zip(meta.rows, values):
        for idx, text in enumerate(row_values):
            row.cells[idx].text = text
    style_table(meta, header=False)

    add_heading(doc, "一、已完成工作")
    finished = [
        f"完成原始知识库整理与清洗：原始 {DATA_PROFILE['raw_rows']} 条，清洗后 {DATA_PROFILE['clean_rows']} 条，文本块 {DATA_PROFILE['chunk_count']} 个。",
        CATEGORY_SUMMARY,
        f"完成 TF-IDF 向量索引与 BM25 关键词索引，TF-IDF 词表 {TRAINING_LOG['vocabulary_size']}，BM25 词表 {TRAINING_LOG['bm25_vocabulary_size']}。",
        f"构建 {QUESTION_COUNT} 条评测问题，完成向量检索、BM25 和混合检索对比；当前选用混合检索，Hit@1={EVAL_SUMMARY['hit_at_1']:.2f}，Hit@3={EVAL_SUMMARY['hit_at_3']:.2f}，MRR={EVAL_SUMMARY['mrr']:.2f}。",
        "完成命令行问答入口、Streamlit/Gradio 演示入口、评估日志和四张可视化图表。",
    ]
    for item in finished:
        add_paragraph(doc, "• " + item, 9.6, space_after=2)

    add_heading(doc, "二、数据处理细节")
    data_table = doc.add_table(rows=6, cols=3)
    rows = [
        ("处理步骤", "具体做法", "输出文件"),
        ("文档抽取", "支持 PDF、DOCX、TXT、MD；真实文档放入 data/source_docs", "extracted_documents.jsonl"),
        ("字段校验", "检查 doc_id、类别、标题、正文、来源、链接、更新时间", "logs/data_profile.json"),
        ("文本清洗", "统一空白字符、去重、过滤过短文本，保留来源引用", "knowledge_base_clean.csv"),
        ("文本切分", "标题与正文拼接，按约260字符切分；后续真实文档按约500字/50字重叠切分", "chunks.csv"),
        ("评测集", f"按校园常见问题标注 gold_doc_id 和关键词，当前 {QUESTION_COUNT} 题，后续扩展至80题以上", "eval_questions.jsonl"),
    ]
    for row, row_values in zip(data_table.rows, rows):
        for idx, text in enumerate(row_values):
            row.cells[idx].text = text
    style_table(data_table)

    add_heading(doc, "三、模型设计思路")
    add_paragraph(
        doc,
        "当前系统采用“问题输入 → 多策略检索 → 证据片段组织 → 带引用回答生成”的流水线。"
        "凌霄负责的检索模块包含三种策略：TF-IDF 向量检索、BM25 关键词检索、0.5/0.5 混合检索。"
        "周子涵负责的生成模块支持通义千问/OpenAI 兼容 API；未配置 Key 时自动使用抽取式回答，保证课堂演示可复现。",
        9.6,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "四、阶段性实验结果")
    result_table = doc.add_table(rows=5, cols=4)
    result_rows = [
        ("指标", "向量检索", "BM25", "混合检索"),
        ("Hit@1", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['hit_at_1']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['hit_at_1']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['hit_at_1']:.2f}"),
        ("Hit@3", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['hit_at_3']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['hit_at_3']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['hit_at_3']:.2f}"),
        ("MRR", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['mrr']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['mrr']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['mrr']:.2f}"),
        ("平均耗时/ms", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['avg_latency_ms']:.3f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['avg_latency_ms']:.3f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['avg_latency_ms']:.3f}"),
    ]
    for row, row_values in zip(result_table.rows, result_rows):
        for idx, text in enumerate(row_values):
            row.cells[idx].text = text
    style_table(result_table)

    figure_path = FIGURE_DIR / "retrieval_strategy_comparison.png"
    if figure_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(figure_path), width=Cm(13.5))

    add_heading(doc, "五、遇到的问题与解决方案")
    issues = doc.add_table(rows=4, cols=3)
    issue_rows = [
        ("问题", "原因分析", "解决方案"),
        ("无 API Key 时无法稳定调用大模型", "课堂环境和网络条件不确定", "保留通义千问/OpenAI 兼容接口，同时实现离线抽取式回答兜底"),
        ("中文短问题检索容易受分词影响", "校园问题表达短且口语化", "采用字符 n-gram TF-IDF，减少分词误差"),
        ("真实资料占比仍需提高", "中期阶段以结构化FAQ跑通闭环", "后续接入学校真实PDF/Word资料，并把评测问题扩展到80题以上"),
    ]
    for row, row_values in zip(issues.rows, issue_rows):
        for idx, text in enumerate(row_values):
            row.cells[idx].text = text
    style_table(issues)

    add_heading(doc, "六、下一步计划")
    next_steps = [
        "补充真实校园手册、学院通知和常见问答，完善数据来源记录。",
        "增加向量模型对比实验：TF-IDF、BGE-small-zh、OpenAI/Qwen Embedding 或 FAISS。",
        "接入通义千问 API，比较 RAG 与无检索基线回答的正确性、引用完整性和用户体验。",
        "完善 Streamlit/Gradio Demo、答辩 PPT、失败案例分析和最终技术报告。",
    ]
    for item in next_steps:
        add_paragraph(doc, "• " + item, 9.6, space_after=2)
    doc.save(path)


def register_pdf_font():
    font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    if not font_path.exists():
        raise FileNotFoundError("Chinese font simhei.ttf not found.")
    pdfmetrics.registerFont(TTFont("SimHei", str(font_path)))


def pdf_styles():
    register_pdf_font()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "title",
            parent=base["Title"],
            fontName="SimHei",
            fontSize=16,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1F4E79"),
            leading=21,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=base["Normal"],
            fontName="SimHei",
            fontSize=12,
            alignment=TA_CENTER,
            leading=16,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "h": ParagraphStyle(
            "h",
            parent=base["Heading2"],
            fontName="SimHei",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#1F4E79"),
            spaceBefore=5,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "body",
            parent=base["BodyText"],
            fontName="SimHei",
            fontSize=9.2,
            leading=13.2,
            alignment=TA_LEFT,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "small": ParagraphStyle(
            "small",
            parent=base["BodyText"],
            fontName="SimHei",
            fontSize=8.2,
            leading=11.2,
            spaceAfter=2,
            wordWrap="CJK",
        ),
    }


def pdf_table(rows, widths=None, font_size=8.2):
    table = Table(rows, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "SimHei"),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("LEADING", (0, 0), (-1, -1), font_size + 3),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9EB6CE")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def build_selection_pdf(path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.45 * cm,
        rightMargin=1.45 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.25 * cm,
    )
    story = [
        Paragraph("《Python 人工智能程序设计实践》选题报告", styles["title"]),
        Paragraph(TITLE, styles["subtitle"]),
        pdf_table(
            [
                ["项目编号", "B07", "提交日期", REPORT_DATE],
                ["项目名称", "RAG 校园智能问答助手", "题目来源", "课程项目库 B 类"],
                ["小组成员", TEAM_MEMBERS, "组队人数", "4人"],
                ["团队分工", TEAM_ASSIGNMENT, "代码仓库", "本地已建，可上传 GitHub/Gitee"],
            ],
            [2.1 * cm, 6.2 * cm, 2.2 * cm, 6.0 * cm],
            8.0,
        ),
        Spacer(1, 0.2 * cm),
        Paragraph("一、项目简介", styles["h"]),
        Paragraph("本项目构建面向校园办事、学习生活与公共服务场景的智能问答助手。系统采用检索增强生成（RAG）范式，先从校园知识库召回相关文本块，再基于证据生成带来源引用的回答，降低通用模型回答不准、无法说明依据的风险，并设置 RAG 与无检索基线、不同检索策略之间的对比实验。", styles["body"]),
        Paragraph("二、数据来源与合规性", styles["h"]),
        Paragraph(f"现阶段已整理模拟校园 FAQ 与办事指南 {DATA_PROFILE['clean_rows']} 条，覆盖 {len(DATA_PROFILE['categories'])} 类场景；字段包括编号、类别、标题、正文、来源、链接和更新时间。后续由归梦依补充学校官网学生手册、教务规定、图书馆与宿舍办事指南等 PDF/Word 原始资料；当前数据不含个人隐私。", styles["body"]),
        Paragraph("三、技术路线", styles["h"]),
        pdf_table([["文档抽取与清洗", "切分、向量化与建库", "BM25/向量/混合检索", "RAG回答、评测与Demo"]], [4.1 * cm] * 4, 8.0),
        Paragraph("基线采用 pdfplumber/python-docx 抽取文档、Pandas 清洗、Scikit-learn TF-IDF 建立向量索引，并实现自定义 BM25 与 0.5/0.5 混合检索。回答模块预留通义千问/OpenAI 兼容 API；无 Key 时使用离线抽取式回答。", styles["body"]),
        Paragraph("四、参考代码与资料", styles["h"]),
        Paragraph("参考 scikit-learn TfidfVectorizer、Streamlit、Gradio、OpenAI Python SDK、DashScope 兼容接口文档，以及 RAG 相关论文和开源实现；实际代码已在本项目目录中独立实现，后续上传仓库时将在 README 中保留引用说明。", styles["body"]),
        Paragraph("五、进度计划", styles["h"]),
        pdf_table(
            [
                ["阶段", "任务", "提交物"],
                ["第7周", "完成选题、数据源确定、项目结构搭建", "选题报告"],
                ["第8-10周", "完成数据清洗、基线检索、BM25/混合检索、评测问题集", "中期进展报告"],
                ["第11-14周", "扩充真实知识库到100条以上，引入 Embedding/FAISS 或大模型 API，完善可视化", "代码仓库更新"],
                ["第15-16周", "整理技术报告、PPT 与演示 Demo，完成答辩", "最终材料"],
            ],
            [2.0 * cm, 11.0 * cm, 3.4 * cm],
            8.0,
        ),
    ]
    doc.build(story)


def build_midterm_pdf(path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.45 * cm,
        rightMargin=1.45 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.25 * cm,
    )
    story = [
        Paragraph("《Python 人工智能程序设计实践》中期进展报告", styles["title"]),
        Paragraph(TITLE, styles["subtitle"]),
        pdf_table(
            [
                ["项目编号", "B07", "提交日期", REPORT_DATE],
                ["小组成员", TEAM_MEMBERS, "当前阶段", "已跑通基线模型"],
                ["代码目录", "当前项目目录", "Demo", "Streamlit + CLI"],
            ],
            [2.1 * cm, 6.2 * cm, 2.2 * cm, 6.0 * cm],
            8.0,
        ),
        Spacer(1, 0.15 * cm),
        Paragraph("一、已完成工作", styles["h"]),
        Paragraph(f"1. 完成原始知识库整理与清洗：原始 {DATA_PROFILE['raw_rows']} 条，清洗后 {DATA_PROFILE['clean_rows']} 条，文本块 {DATA_PROFILE['chunk_count']} 个。", styles["body"]),
        Paragraph(f"2. {CATEGORY_SUMMARY}", styles["body"]),
        Paragraph(f"3. 完成 TF-IDF 向量索引与 BM25 关键词索引，TF-IDF 词表 {TRAINING_LOG['vocabulary_size']}，BM25 词表 {TRAINING_LOG['bm25_vocabulary_size']}。", styles["body"]),
        Paragraph(f"4. 构建 {QUESTION_COUNT} 条评测问题，完成向量检索、BM25 和混合检索对比；当前选用混合检索，Hit@1={EVAL_SUMMARY['hit_at_1']:.2f}，Hit@3={EVAL_SUMMARY['hit_at_3']:.2f}，MRR={EVAL_SUMMARY['mrr']:.2f}。", styles["body"]),
        Paragraph("5. 完成命令行问答入口、Streamlit/Gradio 演示入口、评估日志和四张可视化图表。", styles["body"]),
        Paragraph("二、数据处理细节", styles["h"]),
        pdf_table(
            [
                ["处理步骤", "具体做法", "输出文件"],
                ["文档抽取", "支持 PDF、DOCX、TXT、MD；真实文档放入 data/source_docs", "extracted_documents.jsonl"],
                ["字段校验", "检查 doc_id、类别、标题、正文、来源、链接、更新时间", "logs/data_profile.json"],
                ["文本清洗", "统一空白字符、去重、过滤过短文本，保留来源引用", "knowledge_base_clean.csv"],
                ["文本切分", "标题与正文拼接，按约260字符切分；真实文档按约500字/50字重叠切分", "chunks.csv"],
                ["评测集", f"按校园常见问题标注 gold_doc_id 和关键词，当前 {QUESTION_COUNT} 题，后续扩展至80题以上", "eval_questions.jsonl"],
            ],
            [2.4 * cm, 10.0 * cm, 4.2 * cm],
            7.6,
        ),
        Paragraph("三、模型设计思路", styles["h"]),
        Paragraph("当前系统采用“问题输入 → 多策略检索 → 证据片段组织 → 带引用回答生成”的流水线。检索模块包含 TF-IDF 向量检索、BM25 关键词检索、0.5/0.5 混合检索；生成模块支持通义千问/OpenAI 兼容 API，未配置 Key 时自动使用抽取式回答，保证课堂演示可复现。", styles["body"]),
        PageBreak(),
        Paragraph("四、阶段性实验结果", styles["h"]),
        pdf_table(
            [
                ["指标", "向量检索", "BM25", "混合检索"],
                ["Hit@1", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['hit_at_1']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['hit_at_1']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['hit_at_1']:.2f}"],
                ["Hit@3", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['hit_at_3']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['hit_at_3']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['hit_at_3']:.2f}"],
                ["MRR", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['mrr']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['mrr']:.2f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['mrr']:.2f}"],
                ["平均耗时/ms", f"{EVAL_SUMMARY['retrieval_strategies']['tfidf_vector']['avg_latency_ms']:.3f}", f"{EVAL_SUMMARY['retrieval_strategies']['bm25_keyword']['avg_latency_ms']:.3f}", f"{EVAL_SUMMARY['retrieval_strategies']['hybrid_50_50']['avg_latency_ms']:.3f}"],
            ],
            [3.2 * cm, 4.3 * cm, 4.3 * cm, 4.3 * cm],
            8.0,
        ),
    ]
    figure = FIGURE_DIR / "retrieval_strategy_comparison.png"
    if figure.exists():
        story += [Spacer(1, 0.18 * cm), Image(str(figure), width=14.5 * cm, height=8.16 * cm)]
    story += [
        Paragraph("五、遇到的问题与解决方案", styles["h"]),
        pdf_table(
            [
                ["问题", "原因分析", "解决方案"],
                ["无 API Key 时无法稳定调用大模型", "课堂环境和网络条件不确定", "保留通义千问/OpenAI 兼容接口，同时实现离线抽取式回答兜底"],
                ["中文短问题检索容易受分词影响", "校园问题表达短且口语化", "采用字符 n-gram TF-IDF，减少分词误差"],
                ["真实资料占比仍需提高", "中期阶段以结构化FAQ跑通闭环", "后续接入学校真实PDF/Word资料，并把评测问题扩展到80题以上"],
            ],
            [4.8 * cm, 5.0 * cm, 6.4 * cm],
            7.2,
        ),
        Paragraph("六、下一步计划", styles["h"]),
        Paragraph("补充真实校园手册、学院通知和常见问答；增加 TF-IDF、BGE-small-zh、OpenAI/Qwen Embedding 或 FAISS 对比实验；接入通义千问 API 比较 RAG 与无检索回答；完善 Streamlit/Gradio Demo、答辩 PPT、失败案例分析和最终技术报告。", styles["body"]),
    ]
    doc.build(story)


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    build_selection_docx(REPORT_DIR / "选题报告_B07_RAG校园智能问答助手.docx")
    build_midterm_docx(REPORT_DIR / "中期进展报告_B07_RAG校园智能问答助手.docx")
    build_selection_pdf(REPORT_DIR / "选题报告_B07_RAG校园智能问答助手.pdf")
    build_midterm_pdf(REPORT_DIR / "中期进展报告_B07_RAG校园智能问答助手.pdf")
    print("Reports generated in", REPORT_DIR)


if __name__ == "__main__":
    main()
